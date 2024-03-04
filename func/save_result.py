from docx import Document
from docx.shared import Inches
import os
from win32com.client import Dispatch
import datetime

def daochu_p3_result(path, content, filename):
    document = Document()
    document.add_heading(u'聚类分析报告', 0)
    now_time = datetime.datetime.now().strftime('%F %T')
    document.add_paragraph(now_time)

    document.add_picture(path + '/temp.png', width=Inches(5.0))


    document.add_paragraph(content)

    document.save(path + '/temp.doc')

    word = Dispatch('Word.Application')
    doc = word.Documents.Open(path + '/temp.doc')
    doc.SaveAs(filename[0],17)
    doc.Close()
    word.Quit()
    os.remove(path + '/temp.png')
    os.remove(path + '/temp.doc')

def daochu_p4_result(path, content, filename):

    document = Document()
    document.add_heading(u'预测分析报告', 0)
    now_time = datetime.datetime.now().strftime('%F %T')
    document.add_paragraph(now_time)

    document.add_picture(path + '/temp.png', width=Inches(5.0))

    document.add_paragraph(content)
    document.save(path + '/temp.doc')

    word = Dispatch('Word.Application')
    doc = word.Documents.Open(path + '/temp.doc')
    doc.SaveAs(filename[0],17)
    doc.Close()
    word.Quit()
    os.remove(path + '/temp.png')
    os.remove(path + '/temp.doc')


